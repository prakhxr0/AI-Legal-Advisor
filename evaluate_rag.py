import os
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

from agentic_rag_core import agentic_app  # Fix 1: correct import name

# ---------------- EMBEDDING MODEL ----------------
model = SentenceTransformer("all-MiniLM-L6-v2")


# ---------------- METRICS ----------------
def answer_similarity(ans, gt):
    emb1 = model.encode([ans])
    emb2 = model.encode([gt])
    return float(cosine_similarity(emb1, emb2)[0][0])


def context_precision(retrieved_docs, query):
    if len(retrieved_docs) == 0:
        return 0
    relevant = 0
    for doc in retrieved_docs:
        if query.lower() in doc["content"].lower():  # Fix 2: "content" not "text"
            relevant += 1
    return relevant / len(retrieved_docs)


def entity_recall(retrieved_docs, expected_entities):
    if len(expected_entities) == 0:
        return 0
    found = 0
    for ent in expected_entities:
        for doc in retrieved_docs:
            if ent.lower() in doc["content"].lower():  # Fix 3: "content" not "text"
                found += 1
                break
    return found / len(expected_entities)


def simple_faithfulness(answer, context):
    # basic heuristic (no hallucination check)
    overlap = sum(1 for word in answer.split() if word.lower() in context.lower())
    return overlap / max(len(answer.split()), 1)


def answer_relevancy(query, answer):
    overlap = sum(1 for word in query.split() if word.lower() in answer.lower())
    return overlap / max(len(query.split()), 1)


def context_relevancy(query, context):
    overlap = sum(1 for word in query.split() if word.lower() in context.lower())
    return overlap / max(len(query.split()), 1)


# ---------------- TEST DATA ----------------
test_data = [
    {
        "query": "Can the Central Government extend cantonment provisions to areas beyond the cantonment limits?",
        "ground_truth": "Yes, under Section 352 of the Cantonments Act 2006, the Central Government may by notification extend provisions of Chapters VIII to XV or any rule or bye-law to areas beyond a cantonment and in its vicinity, with or without restrictions or modifications.",
        "entities": ["Central Government", "Cantonments Act", "Official Gazette"]
    },
    {
        "query": "Can a Cantonment Board delegate its functions to the President or Chief Executive Officer?",
        "ground_truth": "Yes, under Section 353 of the Cantonments Act 2006, the Board may by resolution delegate functions to the President, Vice-President, Chief Executive Officer or Health Officer under specified sections, subject to conditions.",
        "entities": ["Cantonment Board", "President", "Chief Executive Officer", "Health Officer", "civil area committee"]
    },
    {
        "query": "Does the Transfer of Property Act apply to cantonments in India?",
        "ground_truth": "Yes, under Section 354 of the Cantonments Act 2006, paragraphs 2 and 3 of Section 54 and Sections 59, 107 and 123 of the Transfer of Property Act 1882 extend to every cantonment from the Act's commencement.",
        "entities": ["Transfer of Property Act", "Registrar", "Sub-Registrar", "Chief Executive Officer", "Defence Estates Officer"]
    },
    {
        "query": "Can a notice issued under the Cantonments Act be invalidated due to a defect of form?",
        "ground_truth": "No, under Section 355 of the Cantonments Act 2006, no notice, order, requisition, licence, or permission in writing issued under the Act shall be invalid merely by reason of any defect of form.",
        "entities": ["Cantonments Act", "notice", "licence"]
    },
    {
        "query": "How can a copy of a document held by a Cantonment Board be admitted as evidence?",
        "ground_truth": "Under Section 356 of the Cantonments Act 2006, a duly certified copy of any receipt, application, plan, notice, order, or register entry in possession of a Board is admissible as evidence of the matters recorded therein.",
        "entities": ["Cantonment Board", "Chief Executive Officer", "Cantonments Act"]
    },
    {
        "query": "Can a Cantonment Board officer be compelled to appear as a witness in legal proceedings?",
        "ground_truth": "Under Section 357 of the Cantonments Act 2006, no officer or employee of a Board is required to produce any register or appear as a witness in proceedings to which the Board is not a party, unless ordered by a court for special cause.",
        "entities": ["Cantonment Board", "court", "Cantonments Act"]
    },
    {
        "query": "How are cantonments treated under the Government Buildings Act 1899?",
        "ground_truth": "Under Section 358 of the Cantonments Act 2006, cantonments and Boards are deemed to be municipalities and municipal authorities respectively for the purposes of the Government Buildings Act 1899, and references to the State Government are read as references to the Central Government.",
        "entities": ["Government Buildings Act", "Central Government", "Cantonment Board", "Cantonments Act"]
    },

    # --- Case Law Test Data (from documents provided earlier) ---
    {
        "query": "What was the Supreme Court's decision in Aparna Bhat vs State of Madhya Pradesh regarding bail conditions?",
        "ground_truth": "The Supreme Court set aside the impugned bail condition imposed by the Madhya Pradesh High Court, expunging it from the record, and disposed of the appeal accordingly.",
        "entities": ["Supreme Court", "Madhya Pradesh High Court", "Aparna Bhat"]
    },
    {
        "query": "Why did the Supreme Court quash the bail granted to the accused in Bohatti Devi vs State of Uttar Pradesh?",
        "ground_truth": "The Supreme Court quashed the bail because the High Court failed to consider the seriousness and gravity of the offence under Sections 302 and 120B IPC, the charge sheet material, and the fact that a non-bailable warrant had already been issued.",
        "entities": ["Supreme Court", "Bohatti Devi", "High Court", "IPC", "Section 302", "Section 120B"]
    },
    {
        "query": "Was the WhatsApp status 'August 5 Black Day Jammu and Kashmir' an offence under Section 153-A IPC?",
        "ground_truth": "No. The Supreme Court in Javed Ahmad Hajam vs State of Maharashtra held that the messages were expressions of protest protected under Article 19(1)(a) and did not promote enmity or hatred between groups as required under Section 153-A IPC.",
        "entities": ["Supreme Court", "Javed Ahmad Hajam", "Section 153-A", "Article 19", "WhatsApp", "Jammu and Kashmir"]
    },
    {
        "query": "What did the Supreme Court rule in Smt. Chintambramma vs State of Karnataka regarding circumstantial evidence?",
        "ground_truth": "The Supreme Court acquitted the appellants holding that the prosecution failed to establish a complete and unbroken chain of circumstances pointing only to the guilt of the accused, and noted the Investigating Officer played a dubious role in the case.",
        "entities": ["Supreme Court", "Smt. Chintambramma", "State of Karnataka", "Section 302", "IPC"]
    },
]


# ---------------- DOCX ----------------
doc = DocxDocument()
doc.add_heading("Agentic RAG Evaluation Report", 0)


# ---------------- RUN EVALUATION ----------------
for i, item in enumerate(test_data, 1):
    query = item["query"]

    state = {
        "user_query": query,
        "sub_issues": [],
        "research_plan": [],
        "retrieved_evidence": [],
        "evaluation_status": "",
        "final_answer": "",
        "citations": [],
        "iteration_count": 0
    }

    result = agentic_app.invoke(state)

    answer = result["final_answer"]
    retrieved = result["retrieved_evidence"]

    context_text = " ".join([d["content"] for d in retrieved])  # Fix 4: "content" not "text"

    # ---- METRICS ----
    metrics = {
        "faithfulness": simple_faithfulness(answer, context_text),
        "answer_relevancy": answer_relevancy(query, answer),
        "context_precision": context_precision(retrieved, query),
        "context_relevancy": context_relevancy(query, context_text),
        "context_entity_recall": entity_recall(retrieved, item["entities"]),
        "answer_similarity": answer_similarity(answer, item["ground_truth"])
    }

    # ---------------- WRITE TO DOCX ----------------
    doc.add_heading(f"Query {i}", level=1)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph(f"Answer: {answer}")
    doc.add_paragraph(f"Citations: {result['citations']}")

    doc.add_heading("Metrics", level=2)
    for k, v in metrics.items():
        doc.add_paragraph(f"{k}: {round(v, 3)}")

    doc.add_paragraph("-" * 50)


# ---------------- SAVE ----------------
doc.save("rag_evaluation_report.docx")

print("Evaluation complete. Saved as rag_evaluation_report.docx")